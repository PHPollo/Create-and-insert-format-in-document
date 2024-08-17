import os

from docx import Document
from docx.shared import Inches

def budget(name_document : str, direction, list_objects : int):
    
    try:
        document = Document(os.path.join(direction, f'{name_document}.docx'))

        #Insertando los elementos del formato.
        document.add_heading(f'{name_document}', 1)

        document.add_paragraph('Descripción del presupuesto')

        document.add_heading('LISTA DE MATERIALES:')


        if list_objects > 0:

            table = document.add_table(rows= list_objects + 2, cols= 4)
            table.style = 'Table Grid'

            
            #table.cell(0, 1).width(Inches(1.5))



            #column = -1

            #while True:
            #    column += 1

            #    for row in table.rows:
            #        cell = row.cells[column]
            #        cell.width = Inches()

            #    break

            table.cell(0, 0).text = 'Materiales' ; table.cell(0, 1).text = 'Unidades (u)' ; table.cell(0, 2).text = '$ Precio Unitario' ; table.cell(0, 3).text = '$ Precio Total'


            table.cell(list_objects + 1, 2). text = '$ Precio Final'

            

            #while True:
                
            #    row += 1
            #    table.cell(row - 1, 0).text = f'{row}° Material' ; table.cell(row - 1, 1).text = '$ Precio'

            #    if row - 1 == list_objects - 1:
            #        table.cell(row, 1).text = '$ Precio Total' ; break
                
        document.save(os.path.join(direction, f'{name_document}.docx'))

        return os.path.join(direction, f'{name_document}.docx')
    except Exception:
        print(f'El documento {name_document} no pudo ser hallado')
        
